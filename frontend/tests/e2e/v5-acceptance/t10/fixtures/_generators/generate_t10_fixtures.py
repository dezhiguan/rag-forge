#!/usr/bin/env python3
"""Generate T10 v5 acceptance fixtures (independent from t10-rewrite assets)."""

from __future__ import annotations

import base64
import io
import struct
import textwrap
import zipfile
import zlib
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
ROOT.mkdir(parents=True, exist_ok=True)


def _png_chunk(kind: bytes, data: bytes) -> bytes:
    return (
        struct.pack(">I", len(data))
        + kind
        + data
        + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF)
    )


def write_png(path: Path, width: int, height: int, pixels: list[tuple[int, int, int]], compress_level: int = 6) -> None:
    raw = bytearray()
    for y in range(height):
        raw.append(0)
        for x in range(width):
            raw.extend(pixels[y * width + x])
    data = b"\x89PNG\r\n\x1a\n"
    data += _png_chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
    data += _png_chunk(b"IDAT", zlib.compress(bytes(raw), compress_level))
    data += _png_chunk(b"IEND", b"")
    path.write_bytes(data)


def geometric_png(path: Path, width: int = 640, height: int = 420, variant: int = 0) -> None:
    pixels: list[tuple[int, int, int]] = []
    for y in range(height):
        for x in range(width):
            bg = 220 - (y * 30 // max(1, height))
            r, g, b = bg, bg + 20, 240 - variant * 8
            if (x - width // 3) ** 2 + (y - height // 2) ** 2 < (min(width, height) // 4) ** 2:
                r, g, b = 255, 180 + variant * 10, 80
            if abs(x - width // 2) < 8:
                r, g, b = 40, 40, 40
            pixels.append((max(0, min(255, r)), max(0, min(255, g)), max(0, min(255, b))))
    write_png(path, width, height, pixels)


def pillow_text_image(path: Path, text: str, subtitle: str = "", variant: int = 0) -> bool:
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return False

    img = Image.new("RGB", (820, 420), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.rectangle((24, 24, 796, 396), outline=(30, 64, 175), width=4)
    draw.rectangle((48, 170, 772, 360), fill=(245, 248, 255))
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 44)
        sub_font = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 28)
    except Exception:
        title_font = ImageFont.load_default()
        sub_font = ImageFont.load_default()
    draw.text((56, 56), text, fill=(15, 23, 42), font=title_font)
    if subtitle:
        draw.text((56, 120), subtitle, fill=(71, 85, 105), font=sub_font)
    draw.text((56, 210), f"variant={variant}", fill=(100, 116, 139), font=sub_font)
    img.save(path)
    return True


def chart_png(path: Path, text: str = "Q1 收入 100w") -> None:
    if pillow_text_image(path, text, "图表截图", variant=1):
        return
    geometric_png(path, variant=3)


def revenue_png(path: Path) -> None:
    if pillow_text_image(path, "今日营收 12345 元", "T10 v5 acceptance OCR fixture"):
        return
    geometric_png(path, variant=5)


def cat_png(path: Path, variant: int) -> None:
    if pillow_text_image(path, "猫咪主题图", f"角度 {variant}", variant=variant):
        return
    geometric_png(path, width=480, height=360, variant=variant)


def corrupt_png(path: Path) -> None:
    path.write_bytes(
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x10\x00\x00\x00\x10\x08\x02\x00\x00\x00"
        + b"\x00" * 64
        + b"CORRUPT"
    )


def oversized_png(path: Path, min_bytes: int = 10_500_000) -> None:
    width, height = 2600, 2600
    pixels: list[tuple[int, int, int]] = []
    for y in range(height):
        for x in range(width):
            pixels.append(((x * 3 + y * 5) % 256, (x + y) % 256, (x * 7 + y) % 256))
    write_png(path, width, height, pixels, compress_level=0)
    if path.stat().st_size < min_bytes:
        with path.open("ab") as f:
            f.write(b"\x00" * (min_bytes - path.stat().st_size))


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _load_rgb(path: Path) -> tuple[int, int, bytes]:
    try:
        from PIL import Image

        img = Image.open(path).convert("RGB")
        return img.width, img.height, img.tobytes()
    except Exception:
        data = path.read_bytes()
        if not data.startswith(b"\x89PNG\r\n\x1a\n"):
            raise ValueError(f"unsupported image format: {path}")
        offset = 8
        width = height = None
        compressed = bytearray()
        while offset < len(data):
            length = struct.unpack(">I", data[offset : offset + 4])[0]
            kind = data[offset + 4 : offset + 8]
            chunk = data[offset + 8 : offset + 8 + length]
            offset += 12 + length
            if kind == b"IHDR":
                width, height, bit_depth, color_type, *_ = struct.unpack(">IIBBBBB", chunk)
                if bit_depth != 8 or color_type != 2:
                    raise ValueError(f"unsupported PNG color type: {path}")
            elif kind == b"IDAT":
                compressed.extend(chunk)
            elif kind == b"IEND":
                break
        if width is None or height is None:
            raise ValueError(f"invalid PNG: {path}")
        raw = zlib.decompress(bytes(compressed))
        stride = width * 3
        rgb = bytearray()
        for y in range(height):
            filter_type = raw[y * (stride + 1)]
            if filter_type != 0:
                raise ValueError(f"unsupported PNG filter {filter_type}: {path}")
            rgb.extend(raw[y * (stride + 1) + 1 : y * (stride + 1) + 1 + stride])
        return width, height, bytes(rgb)


def _pdf_image_object(width: int, height: int, compressed: bytes) -> bytes:
    return (
        (
            f"<< /Type /XObject /Subtype /Image /Width {width} /Height {height} "
            f"/ColorSpace /DeviceRGB /BitsPerComponent 8 /Filter /FlateDecode /Length {len(compressed)} >>\n"
        ).encode()
        + b"stream\n"
        + compressed
        + b"\nendstream"
    )


def _pdf_broken_image_object() -> bytes:
    broken = b"BROKEN_IMAGE_DATA"
    return (
        (
            "<< /Type /XObject /Subtype /Image /Width 16 /Height 16 "
            "/ColorSpace /DeviceRGB /BitsPerComponent 8 /Length "
            f"{len(broken)} >>\n"
        ).encode()
        + b"stream\n"
        + broken
        + b"\nendstream"
    )


def write_multipage_pdf(path: Path, page_specs: list[dict]) -> None:
    objects: list[bytes] = [b"<< /Type /Catalog /Pages 2 0 R >>"]
    page_obj_ids: list[int] = []
    next_id = 4

    objects.append(b"__PAGES_PLACEHOLDER__")
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    for page_no, spec in enumerate(page_specs, start=1):
        text = spec.get("text", f"第{page_no}页")
        image_paths = [Path(p) for p in spec.get("images", [])]
        image_entries: list[tuple[str, int]] = []

        for image_path in image_paths:
            if str(image_path).endswith("mm-corrupt-image.png"):
                image_obj_id = next_id
                next_id += 1
                objects.append(_pdf_broken_image_object())
                image_entries.append((f"Im{page_no}_{len(image_entries)}", image_obj_id))
                continue
            width, height, rgb = _load_rgb(image_path)
            compressed = zlib.compress(rgb, 6)
            image_obj_id = next_id
            next_id += 1
            objects.append(_pdf_image_object(width, height, compressed))
            image_entries.append((f"Im{page_no}_{len(image_entries)}", image_obj_id))

        xobj = ""
        if image_entries:
            xobj = " /XObject << " + " ".join(f"/{name} {obj_id} 0 R" for name, obj_id in image_entries) + " >>"
        resources = f"<< /Font << /F1 3 0 R >>{xobj} >>"

        ops = [f"BT /F1 11 Tf 42 760 Td ({_pdf_escape(text)}) Tj ET"]
        positions = [(42, 520), (320, 520), (42, 250), (320, 250)]
        for idx, (name, _) in enumerate(image_entries):
            x, y = positions[idx % len(positions)]
            ops.append(f"q 220 0 0 110 {x} {y} cm /{name} Do Q")
        stream = "\n".join(ops).encode("latin-1", "ignore")
        content_obj_id = next_id
        next_id += 1
        objects.append(f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream")

        page_obj_id = next_id
        next_id += 1
        page_obj_ids.append(page_obj_id)
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources {resources} /Contents {content_obj_id} 0 R >>"
            ).encode()
        )

    pages_obj = (
        f"<< /Type /Pages /Kids [{' '.join(f'{pid} 0 R' for pid in page_obj_ids)}] /Count {len(page_specs)} >>"
    ).encode()
    objects[1] = pages_obj

    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, 1):
        offsets.append(len(out))
        out.extend(f"{i} 0 obj\n".encode())
        out.extend(obj)
        out.extend(b"\nendobj\n")
    xref = len(out)
    out.extend(f"xref\n0 {len(objects)+1}\n0000000000 65535 f \n".encode())
    for off in offsets[1:]:
        out.extend(f"{off:010d} 00000 n \n".encode())
    out.extend(f"trailer << /Root 1 0 R /Size {len(objects)+1} >>\nstartxref\n{xref}\n%%EOF\n".encode())
    path.write_bytes(bytes(out))


def paragraph_cn(page_no: int) -> str:
    base = (
        "本页为 T10 多模态验收混合 PDF 的第 {page} 页正文段落。"
        "此处包含足够的中文字符用于触发默认 chunker 与 cleaner 管线，"
        "并验证 TEXT chunk 与 IMAGE chunk 能在统一 vl_vector 向量空间共存。"
        "段落还提到知识库检索、DashScope OCR、qwen3-vl-embedding、"
        "以及 PostgreSQL pgvector HNSW 索引等关键组件，确保文本可被稳定召回。"
    )
    return base.format(page=page_no)


def write_docx(path: Path, image_paths: list[Path]) -> None:
    try:
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading("T10 Word 内嵌图验收", level=1)
        doc.add_paragraph("第一段说明：Word 文档应同时产出 TEXT 与 IMAGE chunk。")
        for idx, image_path in enumerate(image_paths, start=1):
            doc.add_paragraph(f"嵌入图 {idx}：")
            doc.add_picture(str(image_path), width=Inches(2.2))
        doc.add_paragraph("第三段总结：三张内嵌图应各自成为 IMAGE chunk。")
        doc.save(path)
        return
    except Exception:
        pass

    # Minimal DOCX zip fallback
    chart = image_paths[0].read_bytes()
    illus = image_paths[1].read_bytes() if len(image_paths) > 1 else chart
    icon = image_paths[2].read_bytes() if len(image_paths) > 2 else chart
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Default Extension="png" ContentType="image/png"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    doc_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image1.png"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image2.png"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image3.png"/>
</Relationships>"""
    document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:body>
    <w:p><w:r><w:t>Word 内嵌图验收：应产出 TEXT 与 IMAGE chunk。</w:t></w:r></w:p>
    <w:p><w:r><w:t>嵌入图 1</w:t></w:r></w:p>
    <w:p><w:r><w:drawing><wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"><a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:blipFill><a:blip r:embed="rId1"/></pic:blipFill></pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>
    <w:p><w:r><w:t>嵌入图 2</w:t></w:r></w:p>
    <w:p><w:r><w:drawing><wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"><a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:blipFill><a:blip r:embed="rId2"/></pic:blipFill></pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>
    <w:p><w:r><w:t>嵌入图 3</w:t></w:r></w:p>
    <w:p><w:r><w:drawing><wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"><a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:blipFill><a:blip r:embed="rId3"/></pic:blipFill></pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>
  </w:body>
</w:document>"""
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/_rels/document.xml.rels", doc_rels)
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/media/image1.png", chart)
        zf.writestr("word/media/image2.png", illus)
        zf.writestr("word/media/image3.png", icon)


def data_uri_png(path: Path) -> str:
    encoded = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:image/png;base64,{encoded}"


def write_html(path: Path, chart_uri: str, illus_uri: str) -> None:
    path.write_text(
        f"""<!DOCTYPE html>
<html lang="zh-CN">
<head><meta charset="UTF-8"><title>T10 HTML img 验收</title></head>
<body>
  <h1>HTML 内嵌图验收</h1>
  <p>本段文本用于产出 TEXT chunk，并验证 HTML img 标签被解析为 IMAGE chunk。</p>
  <img src="{chart_uri}" alt="chart with Q1 revenue">
  <p>中间段落：图表应包含 Q1 收入 100w 等 OCR 文本。</p>
  <img src="{illus_uri}" alt="illustration">
  <p>结尾段落：共两张内嵌图。</p>
</body>
</html>
""",
        encoding="utf-8",
    )


def write_markdown(path: Path, chart_uri: str) -> None:
    path.write_text(
        f"""# T10 Markdown 图片语法验收

本段 Markdown 文本应产出 TEXT chunk。

![Q1 收入图表]({chart_uri})

结尾段落确认 Markdown 图片语法被解析。
""",
        encoding="utf-8",
    )


def main() -> None:
    revenue = ROOT / "mm-pure-image-with-text.png"
    no_text = ROOT / "mm-pure-image-no-text.png"
    chart = ROOT / "_gen-chart.png"
    illus = ROOT / "_gen-illus.png"

    revenue_png(revenue)
    geometric_png(no_text, variant=0)
    chart_png(chart)
    geometric_png(illus, variant=2)

    for idx, name in enumerate(["a", "b", "c", "query"], start=1):
        cat_png(ROOT / f"mm-cat-angle-{name}.png", idx)

    corrupt_png(ROOT / "mm-corrupt-image.png")
    oversized_png(ROOT / "mm-oversized-image.png")

    text_pages = [{"text": paragraph_cn(i)} for i in range(1, 4)]
    write_multipage_pdf(ROOT / "mm-pdf-text-only.pdf", text_pages)

    image_only_pages = [{"text": f"第{i}页纯插图", "images": [illus]} for i in range(1, 4)]
    write_multipage_pdf(ROOT / "mm-pdf-image-only.pdf", image_only_pages)

    mixed_pages = []
    for i in range(1, 6):
        mixed_pages.append(
            {
                "text": paragraph_cn(i),
                "images": [chart, illus],
            }
        )
    write_multipage_pdf(ROOT / "mm-pdf-mixed-rich.pdf", mixed_pages)

    write_multipage_pdf(
        ROOT / "mm-pdf-corrupt-mixed.pdf",
        [
            {
                "text": "损坏图 graceful 验收：含一张正常图与一张损坏图。",
                "images": [chart, ROOT / "mm-corrupt-image.png"],
            }
        ],
    )

    write_docx(ROOT / "mm-word-with-embedded-images.docx", [chart, illus, revenue])

    chart_uri = data_uri_png(chart)
    illus_uri = data_uri_png(illus)
    write_html(ROOT / "mm-html-img-tag.html", chart_uri, illus_uri)
    write_markdown(ROOT / "mm-markdown-img-syntax.md", chart_uri)

    print(f"fixtures generated under {ROOT}")


if __name__ == "__main__":
    main()

# -*- coding: utf-8 -*-
"""生成「平台基准库·v1」全格式真实语料（RAGForge 平台技术知识，事实取自 CLAUDE.md，无 PII）。
覆盖 10 类格式 + ZIP/TAR.GZ 归档。所有内容为有意义真实事实，禁填充。"""
import os, csv, zipfile, tarfile
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas as pdfcanvas
from docx import Document
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont

OUT = os.path.join(os.path.dirname(__file__), "corpus")
os.makedirs(OUT, exist_ok=True)
CJK = "/System/Library/Fonts/PingFang.ttc"

def w(name, text):
    with open(os.path.join(OUT, name), "w", encoding="utf-8") as f:
        f.write(text)
    print("  写出", name)

# ---------- 1. Markdown：架构总览 ----------
w("01_架构总览.md", """# RAGForge 平台架构总览

RAGForge 是一个 RAG（检索增强生成）知识引擎，定位为基础设施层，为 CareerMate 等上层应用提供知识检索、RAG 应答与 MCP API 能力。线上域名为 ragforge.net。

## 技术栈
- 后端采用 Java 21 与 Spring Boot 3.5.15，持久层使用 MyBatis-Plus。
- 前端采用 Vue 3 与 Vite，使用纯 JavaScript 编写。

## 检索策略
平台提供五种检索策略，统一由 RetrievalService 调度：
1. vector：默认策略，基于向量相似度。
2. keyword：基于 Elasticsearch 的 BM25 关键词检索。
3. hybrid：向量与关键词的 RRF 融合。
4. rewrite：先改写查询再走多路向量。
5. full：改写 + 混合 + 重排，是唯一调用 Rerank 的策略，默认并发为 1。

## 向量维度
文档向量统一为 2560 维。由于 2560 维超过了 pgvector 0.8 的索引维度上限 2000，当前向量检索没有 HNSW 索引，走的是顺序扫描。
""")

# ---------- 2. TXT：常见问题 FAQ ----------
w("05_常见问题FAQ.txt", """RAGForge 常见问题（FAQ）

问：向量检索为什么走顺序扫描？
答：因为文档向量是 2560 维，超过了 pgvector 0.8 的索引上限 2000 维，无法建立 HNSW 索引，所以当前走顺序扫描。

问：关键词检索依赖什么？
答：依赖 Elasticsearch 8.15 的 BM25。中文分词使用 IK 插件；当 IK 插件缺失时，回退到 standard 分词器。

问：MCP 接口基于什么框架？
答：基于 Spring AI 的 MCP WebMVC SSE 实现。

问：文档处理管道用什么消息队列？
答：使用 RocketMQ，主题为 ragforge-document-process，消费组为 ragforge-doc-process-group。

问：认证撤销与限流依赖什么？
答：依赖 Redis，用于认证撤销、API Key 限流与 ShedLock 分布式锁。
""")

# ---------- 3. HTML：模型与成本说明 ----------
w("04_模型与成本说明.html", """<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8">
<title>RAGForge 模型与成本说明</title></head><body>
<h1>RAGForge 模型与成本说明</h1>
<p>RAGForge 按用途（Purpose）选择模型，多数走阿里云 DashScope，评测判分走 DeepSeek。</p>
<table border="1"><thead><tr><th>用途</th><th>模型</th><th>供应商</th><th>说明</th></tr></thead>
<tbody>
<tr><td>EMBEDDING</td><td>qwen3-vl-embedding</td><td>DashScope</td><td>文本与图片统一 2560 维</td></tr>
<tr><td>REWRITE</td><td>qwen-turbo</td><td>DashScope</td><td>查询改写，支持动态选型</td></tr>
<tr><td>ANSWER</td><td>qwen-plus</td><td>DashScope</td><td>RAG 应答与调试台</td></tr>
<tr><td>RERANK</td><td>qwen3-rerank</td><td>DashScope</td><td>仅 full 策略调用</td></tr>
<tr><td>OCR</td><td>qwen-vl-ocr</td><td>DashScope</td><td>图片管道可选</td></tr>
<tr><td>JUDGE</td><td>deepseek-v4-flash</td><td>DeepSeek</td><td>LLM-as-Judge 评测判分</td></tr>
</tbody></table>
<p>评测判分（JUDGE）的成本归属到系统组织。检索与应答的模型用量按发起组织计量。</p>
</body></html>""")

# ---------- 4. CSV：检索策略指标 ----------
with open(os.path.join(OUT, "06_检索策略指标.csv"), "w", encoding="utf-8-sig", newline="") as f:
    wr = csv.writer(f)
    wr.writerow(["策略", "是否调用重排", "默认并发", "向量维度", "说明"])
    wr.writerow(["vector", "否", "48", "2560", "默认策略，向量相似度"])
    wr.writerow(["keyword", "否", "32", "-", "Elasticsearch BM25 关键词"])
    wr.writerow(["hybrid", "否", "20", "2560", "向量与关键词 RRF 融合"])
    wr.writerow(["rewrite", "否", "16", "2560", "改写查询后多路向量"])
    wr.writerow(["full", "是", "1", "2560", "改写+混合+重排，唯一调用 Rerank"])
    print("  写出 06_检索策略指标.csv")

# ---------- 5. PDF：部署运维手册 ----------
pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
pdf = os.path.join(OUT, "02_部署运维手册.pdf")
c = pdfcanvas.Canvas(pdf, pagesize=A4)
W, H = A4
def line(x, y, s, size=12, bold=False):
    c.setFont("STSong-Light", size)
    c.drawString(x, y, s)
y = H - 30*mm
for s, sz in [
    ("RAGForge 部署运维手册", 18),
    ("", 6),
    ("一、部署形态", 14),
    ("RAGForge 应用层部署在 k3s 单节点集群，命名空间为 ragforge。", 12),
    ("同一个 backend 镜像按环境变量 RAGFORGE_ROLE 启动为不同角色：", 12),
    ("  · api：3 个副本，对外提供 REST 接口。", 12),
    ("  · worker：1 个副本，消费 RocketMQ 处理文档。", 12),
    ("  · judge：1 个副本，运行 LLM-as-Judge 评测判分。", 12),
    ("  · frontend：2 个副本，提供前端页面。", 12),
    ("", 6),
    ("二、入口与端口", 14),
    ("请求路径为：域名 → 入口层 Nginx → 应用层 NodePort 31090。", 12),
    ("", 6),
    ("三、文件存储", 14),
    ("文件存储使用节点本地 hostPath，挂载路径为 /data/files。", 12),
    ("对象存储抽象已就绪，但默认使用本地盘。", 12),
    ("", 6),
    ("四、数据层", 14),
    ("数据层独立部署，包含 PostgreSQL（含 pgvector）、Elasticsearch、RocketMQ 与 Redis。", 12),
]:
    if s == "":
        y -= sz*mm/3
        continue
    line(25*mm, y, s, sz)
    y -= (sz+4)
c.showPage(); c.save()
print("  写出 02_部署运维手册.pdf")

# ---------- 6. DOCX：安全与权限规范 ----------
doc = Document()
doc.add_heading("RAGForge 安全与权限规范", level=0)
doc.add_heading("一、认证机制", level=1)
doc.add_paragraph("认证由 Auth Gateway 颁发 JWT，签名算法为 RS256。后端使用自研的 JwtVerifier，通过 JWKS 验签，而非 nimbus 库。")
doc.add_heading("二、角色模型", level=1)
doc.add_paragraph("平台角色为字符串约定，包含 ADMIN、KB_EDITOR、KB_VIEWER 与 SERVICE_ACCOUNT 四种。知识库访问统一经过 KbAccessGuard 校验。")
doc.add_heading("三、会话策略", level=1)
doc.add_paragraph("access token 有效期为 15 分钟。refresh token 采用旋转机制，有效期 7 天；开启记住我后为 30 天滑动续期。网关提供 60 秒的旋转宽限期。")
doc.add_heading("四、组织模型", level=1)
doc.add_paragraph("组织模型为 GitHub 式的个人加组织结构，已移除早期的 tenant 概念。系统组织的 org_id 为 0，绑定平台超级管理员。")
doc.save(os.path.join(OUT, "03_安全与权限规范.docx"))
print("  写出 03_安全与权限规范.docx")

# ---------- 7. 图片：PNG/JPG/GIF/WEBP（含真实中文文字，供 OCR）----------
def make_image(name, title, lines, size=(1000, 620), fmt=None):
    img = Image.new("RGB", size, (255, 255, 255))
    d = ImageDraw.Draw(img)
    ft = ImageFont.truetype(CJK, 40)
    fh = ImageFont.truetype(CJK, 30)
    d.rectangle([0, 0, size[0], 78], fill=(37, 99, 235))
    d.text((30, 18), title, font=ft, fill=(255, 255, 255))
    y = 115
    for ln in lines:
        d.text((40, y), ln, font=fh, fill=(15, 31, 61))
        y += 52
    img.save(os.path.join(OUT, name), format=fmt)
    print("  写出", name)

make_image("07_架构分层图.png", "RAGForge 架构分层",
           ["入口层：入口层 Nginx 反向代理",
            "应用层：k3s 集群，NodePort 31090",
            "数据层：PostgreSQL + Elasticsearch",
            "消息与缓存：RocketMQ + Redis",
            "向量维度：2560 维，顺序扫描"], fmt="PNG")

make_image("08_部署角色卡.jpg", "k3s 部署角色副本数",
           ["api 服务：3 个副本",
            "worker 服务：1 个副本",
            "judge 服务：1 个副本",
            "frontend 前端：2 个副本",
            "入口端口：NodePort 31090"], fmt="JPEG")

make_image("09_检索策略图.gif", "五种检索策略",
           ["vector：默认，向量相似度",
            "keyword：BM25 关键词",
            "hybrid：RRF 融合",
            "rewrite：改写后多路向量",
            "full：改写+混合+重排，调用 Rerank"], fmt="GIF")

make_image("10_关键指标卡.webp", "平台关键指标",
           ["向量维度：2560 维",
            "pgvector 索引上限：2000 维",
            "当前索引：无 HNSW，顺序扫描",
            "access token 有效期：15 分钟",
            "记住我时长：30 天滑动"], fmt="WEBP")

# ---------- 8. 打包 ZIP 与 TAR.GZ（验证归档上传展开链路）----------
files = sorted(os.listdir(OUT))
files = [f for f in files if not f.endswith((".zip", ".tar.gz"))]
zip_path = os.path.join(OUT, "..", "platform-baseline-corpus.zip")
with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
    for f in files:
        z.write(os.path.join(OUT, f), f)
tgz_path = os.path.join(OUT, "..", "platform-baseline-corpus.tar.gz")
with tarfile.open(tgz_path, "w:gz") as t:
    for f in files:
        t.add(os.path.join(OUT, f), arcname=f)
print("打包完成：", os.path.basename(zip_path), "/", os.path.basename(tgz_path))
print("语料文件数：", len(files))

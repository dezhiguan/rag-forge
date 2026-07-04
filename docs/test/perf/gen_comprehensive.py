#!/usr/bin/env python3
"""全面多格式+多模态语料。每文档独立主题+查询,用于覆盖入库与检索质量评估。"""
import os, csv, json, zipfile
from PIL import Image, ImageDraw, ImageFont
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import docx

OUT = "ccorpus"; os.makedirs(OUT, exist_ok=True)
pdfmetrics.registerFont(TTFont("cjk", "/Library/Fonts/Arial Unicode.ttf"))
PIL_FONT = "/System/Library/Fonts/PingFang.ttc"
manifest = []
def add(fn, ctype, fmt, modality, query): manifest.append({"file":fn,"ctype":ctype,"format":fmt,"modality":modality,"query":query})

def txt(name, content, q): open(f"{OUT}/{name}.txt","w").write(content); add(f"{name}.txt","text/plain","txt","text",q)
def md(name, content, q): open(f"{OUT}/{name}.md","w").write(content); add(f"{name}.md","text/markdown","md","text",q)
def html(name, content, q): open(f"{OUT}/{name}.html","w").write(content); add(f"{name}.html","text/html","html","text",q)
def csvf(name, rows, q):
    with open(f"{OUT}/{name}.csv","w",newline="") as f: csv.writer(f).writerows(rows)
    add(f"{name}.csv","text/csv","csv","text",q)
def pdf(name, title, paras, q, img=None, modality="text"):
    c=canvas.Canvas(f"{OUT}/{name}.pdf",pagesize=A4); c.setFont("cjk",16); c.drawString(70,780,title); c.setFont("cjk",12); y=740
    for p in paras: c.drawString(70,y,p); y-=26
    if img: c.drawImage(img,70,y-220,width=360,height=200)
    c.save(); add(f"{name}.pdf","application/pdf","pdf",modality,q)
def wordf(name, title, paras, q):
    d=docx.Document(); d.add_heading(title,level=1)
    for p in paras: d.add_paragraph(p)
    d.save(f"{OUT}/{name}.docx"); add(f"{name}.docx","application/vnd.openxmlformats-officedocument.wordprocessingml.document","docx","text",q)
def render(path, lines, w=800, h=460):
    im=Image.new("RGB",(w,h),(250,250,252)); dr=ImageDraw.Draw(im)
    tf=ImageFont.truetype(PIL_FONT,34); bf=ImageFont.truetype(PIL_FONT,26)
    dr.rectangle([20,20,w-20,h-20],outline=(60,110,240),width=3); dr.text((45,45),lines[0],fill=(20,30,60),font=tf); y=125
    for ln in lines[1:]: dr.text((45,y),ln,fill=(40,50,80),font=bf); y+=44
    im.save(path)
def imgpdf(name, lines, q):  # 纯图片PDF(图片为主体)
    p=f"{OUT}/_{name}.png"; render(p,lines); pdf(name,"", [], q, img=p, modality="image")
    # 覆盖:纯图片PDF去掉标题行
    c=canvas.Canvas(f"{OUT}/{name}.pdf",pagesize=A4); c.drawImage(p,50,420,width=500,height=300); c.save()

# 纯文本类
txt("asyncio","Python 的 asyncio 通过事件循环调度协程,await 挂起当前协程让出控制权,IO 就绪后事件循环恢复执行,单线程实现高并发。","Python 异步协程怎么调度")
txt("nginx","Nginx 采用 epoll 事件驱动异步非阻塞,master-worker 多进程,单机支撑数万并发连接,常作反向代理负载均衡。","Nginx 为什么能扛高并发")
txt("tcp","TCP 建立连接三次握手 SYN/SYN-ACK/ACK,断开四次挥手,确保双向数据传输完毕保证可靠。","TCP 三次握手过程")
md("kafka","# Kafka\nKafka 用**分区**实现并行消费,消费者组内每分区只被一个消费者消费,offset 记录消费位置保证顺序高吞吐。","Kafka 怎么并行消费")
md("raft","# Raft\nRaft 通过 **leader 选举**和**日志复制**保证一致性,多数派确认后提交,比 Paxos 易懂。","Raft 一致性怎么保证")
csvf("db_cmp",[["数据库","索引","场景"],["MySQL","B+树聚簇","事务OLTP"],["Qdrant","HNSW图","向量检索"],["Redis","跳表","缓存"]],"MySQL 用什么索引结构")
csvf("http_code",[["码","含义"],["200","成功"],["404","资源不存在"],["429","限流"],["503","服务不可用"]],"HTTP 429 是什么意思")
html("grpc","<html><body><h1>gRPC</h1><p>gRPC 基于 HTTP/2 和 Protobuf,支持四种流模式,强类型契约,性能高于 REST JSON。</p></body></html>","gRPC 基于什么协议")
html("cdn","<html><body><h1>CDN</h1><p>CDN 把内容缓存到边缘节点,用户就近访问降低延迟,未命中回源拉取。</p></body></html>","CDN 加速原理")
pdf("mvcc","PostgreSQL MVCC",["PostgreSQL 用 MVCC 多版本并发控制,每行带 xmin/xmax 事务号。","读不阻塞写,VACUUM 回收死元组释放空间。"],"PostgreSQL MVCC 原理")
pdf("jvm_g1","JVM G1 回收器",["G1 把堆分成 Region 区域,并发标记后按回收价值优先清理。","可控制每次停顿时间目标。"],"JVM G1 垃圾回收怎么工作")
wordf("redis_persist","Redis 持久化",["Redis 提供 RDB 快照和 AOF 日志两种持久化。","RDB 定期全量,AOF 记录每条写命令,生产常混合使用。"],"Redis 有哪几种持久化方式")
wordf("oauth2","OAuth2 授权码",["OAuth2 授权码模式:客户端拿授权码换 access token。","token 过期用 refresh token 刷新。"],"OAuth2 授权码流程")
# 纯图片(png/jpg,考验 OCR+VL)
render(f"{OUT}/hnsw.png",["向量数据库 HNSW 索引","分层可导航小世界图","近似最近邻,复杂度接近对数级"]); add("hnsw.png","image/png","png","image","向量数据库 HNSW 图索引原理")
render(f"{OUT}/prometheus.jpg",["Prometheus 监控","Pull 模式抓取指标","PromQL 查询 Grafana 可视化"]); Image.open(f"{OUT}/prometheus.jpg").convert("RGB").save(f"{OUT}/prometheus.jpg","JPEG"); add("prometheus.jpg","image/jpeg","jpg","image","Prometheus 怎么采集监控指标")
# 图片PDF
imgpdf("docker_img",["Docker 镜像分层","每层只读,容器加可写层","copy-on-write 写时复制"],"Docker 镜像为什么分层")
# 图文混合PDF(文字+内嵌图)
render(f"{OUT}/_k8sfig.png",["Kubernetes Pod 调度","亲和性/反亲和性","按节点标签调度"])
pdf("k8s_mix","Kubernetes 调度",["Kubernetes 用亲和性控制 Pod 调度到指定节点。","下图展示调度关系。"],"k8s 怎么控制 pod 调度到哪个节点",img=f"{OUT}/_k8sfig.png",modality="mixed")
# zip 压缩包(含3个txt)
zf=f"{OUT}/bundle.zip"
with zipfile.ZipFile(zf,"w") as z:
    z.writestr("elasticsearch.txt","Elasticsearch 倒排索引把词映射到文档列表,BM25 按词频和逆文档频率打分排序。")
    z.writestr("rabbitmq.txt","RabbitMQ 用 Exchange 路由消息到队列,有 direct/topic/fanout 类型,支持 ack 确认。")
    z.writestr("mysql_innodb.txt","MySQL InnoDB 用 B+ 树聚簇索引,叶子存整行,二级索引存主键再回表。")
add("bundle.zip","application/zip","zip","text","Elasticsearch 全文检索的打分算法")

json.dump(manifest,open(f"{OUT}/manifest.json","w"),ensure_ascii=False,indent=1)
from collections import Counter
print(f"生成 {len(manifest)} 个文档")
print("格式:",dict(Counter(m['format'] for m in manifest)))
print("模态:",dict(Counter(m['modality'] for m in manifest)))
